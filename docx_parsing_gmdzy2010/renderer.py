import os
import re
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt
from docx.oxml.ns import qn

from docx_parsing_gmdzy2010.elements import Elements, element_splitting
from docx_parsing_gmdzy2010.settings import PARAGRAPH_FORMAT, ALL_STYLES


class DocxProduce:
    """The paragraph context, table context or picture context MUST be ready
    to participating into the process of template rendering. Optionally, the
    same to the paragraph format and styles."""
    picture_path = os.path.dirname(os.path.abspath(__file__))
    
    def __init__(self, template_text=None, template_docx=None,
                 para_format=PARAGRAPH_FORMAT, all_styles=ALL_STYLES,
                 paragraph_context=None, table_context=None,
                 picture_context=None):
        self.document = Document(docx=template_docx)
        self.template_docx = template_docx
        self.template_text = template_text
        self.context = paragraph_context
        self.para_format = para_format
        self.table_context = table_context
        self.picture_context = picture_context
        self.contents = None
        self.styles = self.document.styles
        self.all_styles = all_styles

    def _add_style(self, style_name, style_type, font_size, base_style=None):
        style = self.styles.add_style(style_name, style_type)
        style.base_style = base_style
        style.font.size = font_size
        return style
    
    def _set_init_style(self):
        """By default, the font type of chinese and western language are Song
        and Times New Roman, respectively."""
        init = self.styles.add_style('init', WD_STYLE_TYPE.PARAGRAPH)
        self.styles['init'].font.name = 'Times New Roman'
        self.styles['init']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        return init
    
    def set_styles(self):
        init = self._set_init_style()
        for name, prop in self.all_styles.items():
            self._add_style(name, prop[0], prop[1], base_style=init)
    
    def get_document(self):
        """Each type of element in text template should mapping to
        corresponding method."""
        self.contents = self.get_contents()
        for content in self.contents:
            if content["element_type"] == "table":
                self.add_table(content)
            elif content["element_type"] == "picture":
                self.add_picture(content)
            else:
                self.add_paragraph(content)
        return self.document
    
    def add_paragraph(self, content):
        """Method to handle paragraph and run"""
        paragraph = self.document.add_paragraph(style=content["element_type"])
        p_format = paragraph.paragraph_format
        format_map = self.para_format.get(content["format_name"])
        p_format.alignment = format_map.get("alignment")
        p_format.line_spacing = format_map.get("line_spacing", 1.5)
        p_format.space_before = format_map.get("space_before")
        p_format.space_after = format_map.get("space_after")
        p_format.page_break_before = format_map.get("page_break_before", False)
        p_format.first_line_indent = format_map.get("first_line_indent")
        for run in content["contents"]:
            actual_run = paragraph.add_run(text=run.get("run_text"))
            actual_run.font.underline = run.get("underline")
            actual_run.font.bold = run.get("bold")
        return paragraph
    
    @staticmethod
    def set_column_content(row_cells, column, row_set):
        for col_index in range(column):
            row_cells[col_index].text = row_set[col_index]
        return row_cells
    
    def add_table(self, content):
        """Method to handle table"""
        # TODO: attributes of rows and cols can automatically gained.
        context = self.table_context.get(content["format_name"])
        column = context["attr"].get("cols")
        table = self.document.add_table(**context["attr"])
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        for index, content_data in enumerate(context["data"]):
            cells = table.rows[index].cells
            table.rows[index].height = Pt(12)
            self.set_column_content(cells, column, content_data)
        return table
    
    def add_picture(self, content):
        # TODO: the unittest is not passed.
        path = self.picture_path
        context = self.picture_context.get(content["format_name"])
        return self.document.add_picture(path, **context)
    
    def get_contents(self, encoding="UTF-8"):
        """This method takes the text template into file object, then generate
        the basic paragraph elements. the elements is a generator, popping out
        every two line."""
        file_obj = open(self.template_text, encoding=encoding)
        elements = Elements(file_obj)
        contents = [
            element_splitting(element, self.context) for element in elements
        ]
        file_obj.close()
        return contents
    
    def save(self, to_path="", file_name="default"):
        self.set_styles()
        self.get_document()
        return self.document.save("{}{}.docx".format(to_path, file_name))


class ContextFields:
    """Class to get all fields of context"""
    
    def __init__(self, template_text=None, encoding="UTF-8"):
        self.template_text = template_text
        self.fields = self.get_fields(encoding=encoding)
    
    def get_fields(self, encoding="UTF-8"):
        """This method return the field list of the input text template."""
        all_fields = []
        file_obj = open(self.template_text, encoding=encoding)
        elements = Elements(file_obj)
        for element in elements:
            *_, paragraph_text = element
            sub_fields = re.findall('[{](.*?)[}]', paragraph_text)
            all_fields.extend(sub_fields if sub_fields else [])
        file_obj.close()
        all_fields = list(set(all_fields))
        return all_fields
